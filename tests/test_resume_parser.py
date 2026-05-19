from pathlib import Path
import subprocess

from docx import Document

from app import resume_parser


def test_extract_docx_reads_paragraphs_tables_and_header(tmp_path):
    file_path = tmp_path / "resume.docx"
    document = Document()
    document.add_paragraph("Иван Иванов")
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Python"
    table.rows[0].cells[1].text = "SQL"
    document.sections[0].header.paragraphs[0].text = "Header note"
    document.save(file_path)

    extracted = resume_parser.extract_text(file_path)

    assert "Иван Иванов" in extracted
    assert "Python" in extracted
    assert "SQL" in extracted
    assert "Header note" in extracted


def test_extract_doc_uses_libreoffice_conversion(tmp_path, monkeypatch):
    source_path = tmp_path / "legacy.doc"
    source_path.write_bytes(b"legacy doc placeholder")

    def fake_run(command, capture_output, text, check):
        output_dir = Path(command[command.index("--outdir") + 1])
        converted = output_dir / "legacy.docx"
        document = Document()
        document.add_paragraph("Converted from DOC")
        document.save(converted)
        return subprocess.CompletedProcess(command, 0, stdout="ok", stderr="")

    monkeypatch.setattr(resume_parser, "_find_soffice_path", lambda: Path("/fake/soffice"))
    monkeypatch.setattr(resume_parser.subprocess, "run", fake_run)

    extracted = resume_parser.extract_text(source_path)

    assert "Converted from DOC" in extracted


def test_get_soffice_status_reports_env_path(tmp_path, monkeypatch):
    soffice = tmp_path / "soffice"
    soffice.write_text("fake", encoding="utf-8")
    monkeypatch.setenv("SOFFICE_PATH", str(soffice))

    status = resume_parser.get_soffice_status()

    assert status == {"available": True, "path": str(soffice)}


def test_extract_doc_reports_missing_libreoffice(tmp_path, monkeypatch):
    source_path = tmp_path / "legacy.doc"
    source_path.write_bytes(b"legacy doc placeholder")
    monkeypatch.setattr(resume_parser, "_find_soffice_path", lambda: None)

    try:
        resume_parser.extract_text(source_path)
    except RuntimeError as exc:
        assert "LibreOffice" in str(exc)
    else:
        raise AssertionError("Expected RuntimeError when LibreOffice is missing")


def test_filename_only_profile_uses_doc_filename_when_libreoffice_missing():
    profile = resume_parser.parse_resume_profile("", "Юрченко Анна Владимировна.doc")

    assert profile["full_name_original"] == "Юрченко Анна Владимировна"
