import os
import re
import shutil
import subprocess
import tempfile
from pathlib import Path

import fitz
from docx import Document

from app.skills import extract_skills, flatten_skills
from app.text_utils import normalize_phone


EMAIL_RE = re.compile(r"[\w.+-]+@[\w-]+\.[\w.-]+", re.IGNORECASE)
PHONE_RE = re.compile(r"(?:\+?\d[\s()\-]*){9,16}")


def extract_text(file_path: Path) -> str:
    suffix = file_path.suffix.lower()
    if suffix == ".pdf":
        return _extract_pdf(file_path)
    if suffix == ".docx":
        return _extract_docx(file_path)
    if suffix == ".doc":
        return _extract_doc(file_path)
    if suffix == ".txt":
        return file_path.read_text(encoding="utf-8", errors="ignore")
    raise ValueError(f"Unsupported resume format: {suffix}")


def parse_resume_profile(text: str, filename: str = "") -> dict[str, object]:
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    skills = extract_skills(text)
    email = _first_match(EMAIL_RE, text)
    phone_raw = _first_match(PHONE_RE, text)
    full_name = _guess_full_name(lines, filename)
    return {
        "full_name_original": full_name,
        "email": email,
        "phone": normalize_phone(phone_raw),
        "city": _field_after_label(text, ["city", "город", "location"]),
        "current_position": _guess_position(lines),
        "current_company": _field_after_label(text, ["company", "компания", "работодатель"]),
        "years_experience": _guess_years_experience(text),
        "education": _field_after_label(text, ["education", "образование"]),
        "english_level": _guess_english_level(text),
        "programming_languages": skills["programming_languages"],
        "embedded_stack": skills["embedded_stack"],
        "protocols": skills["protocols"],
        "tools": skills["tools"],
        "key_skills": flatten_skills(skills),
    }


def empty_resume_profile(filename: str = "") -> dict[str, object]:
    return {
        "full_name_original": "",
        "email": "",
        "phone": "",
        "city": "",
        "current_position": "",
        "current_company": "",
        "years_experience": "",
        "education": "",
        "english_level": "",
        "programming_languages": [],
        "embedded_stack": [],
        "protocols": [],
        "tools": [],
        "key_skills": [],
    }


def _extract_pdf(file_path: Path) -> str:
    parts = []
    with fitz.open(file_path) as doc:
        for page in doc:
            parts.append(page.get_text())
    return "\n".join(parts)


def _extract_docx(file_path: Path) -> str:
    document = Document(file_path)
    parts: list[str] = []
    for paragraph in document.paragraphs:
        _append_text(parts, paragraph.text)
    for table in document.tables:
        _append_table(parts, table)
    for section in document.sections:
        for paragraph in section.header.paragraphs:
            _append_text(parts, paragraph.text)
        for table in section.header.tables:
            _append_table(parts, table)
        for paragraph in section.footer.paragraphs:
            _append_text(parts, paragraph.text)
        for table in section.footer.tables:
            _append_table(parts, table)
    return "\n".join(parts)


def _extract_doc(file_path: Path) -> str:
    with tempfile.TemporaryDirectory() as tmp_dir:
        converted = _convert_doc_to_docx(file_path, Path(tmp_dir))
        return _extract_docx(converted)


def _first_match(pattern: re.Pattern[str], text: str) -> str:
    match = pattern.search(text)
    return match.group(0).strip() if match else ""


def _guess_full_name(lines: list[str], filename: str) -> str:
    for line in lines[:8]:
        cleaned = re.sub(r"\s+", " ", line).strip()
        words = cleaned.split()
        if 2 <= len(words) <= 4 and not EMAIL_RE.search(cleaned) and not PHONE_RE.search(cleaned):
            if not re.search(r"(resume|cv|резюме|engineer|developer|manager)", cleaned, re.IGNORECASE):
                return cleaned
    return Path(filename).stem.replace("_", " ").replace("-", " ")


def _guess_position(lines: list[str]) -> str:
    patterns = ["engineer", "developer", "manager", "architect", "разработчик", "инженер", "руководитель"]
    for line in lines[:20]:
        if any(pattern in line.lower() for pattern in patterns):
            return line[:160]
    return ""


def _field_after_label(text: str, labels: list[str]) -> str:
    for label in labels:
        match = re.search(rf"{label}\s*[:\-]\s*(.+)", text, re.IGNORECASE)
        if match:
            return match.group(1).strip()[:160]
    return ""


def _guess_years_experience(text: str) -> str:
    patterns = [
        r"(\d{1,2})\+?\s*(?:years|yrs|лет|года)\s+(?:experience|опыта)",
        r"(?:experience|опыт)\D{0,20}(\d{1,2})\+?",
    ]
    for pattern in patterns:
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            return match.group(1)
    return ""


def _guess_english_level(text: str) -> str:
    match = re.search(r"\b(A1|A2|B1|B2|C1|C2|Intermediate|Upper Intermediate|Advanced|Fluent)\b", text, re.IGNORECASE)
    return match.group(0) if match else ""


def _append_table(parts: list[str], table) -> None:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                _append_text(parts, paragraph.text)
            for nested_table in cell.tables:
                _append_table(parts, nested_table)


def _append_text(parts: list[str], text: str) -> None:
    cleaned = str(text or "").strip()
    if cleaned:
        parts.append(cleaned)


def _convert_doc_to_docx(file_path: Path, output_dir: Path) -> Path:
    soffice_path = _find_soffice_path()
    if not soffice_path:
        raise RuntimeError("Для файлов .doc требуется установленный LibreOffice")

    command = [
        str(soffice_path),
        "--headless",
        "--convert-to",
        "docx",
        "--outdir",
        str(output_dir),
        str(file_path),
    ]
    try:
        completed = subprocess.run(command, capture_output=True, text=True, check=True)
    except subprocess.CalledProcessError as exc:
        message = exc.stderr.strip() or exc.stdout.strip() or "неизвестная ошибка конвертации"
        raise RuntimeError(f"Не удалось конвертировать .doc через LibreOffice: {message[:240]}") from exc

    converted = output_dir / f"{file_path.stem}.docx"
    if converted.exists():
        return converted

    candidates = sorted(output_dir.glob("*.docx"))
    if candidates:
        return candidates[0]

    output = (completed.stdout or completed.stderr or "").strip()
    raise RuntimeError(f"LibreOffice не создал .docx-файл: {output[:240]}")


def _find_soffice_path() -> Path | None:
    env_candidates = [
        os.environ.get("LIBREOFFICE_PATH"),
        os.environ.get("SOFFICE_PATH"),
    ]
    for candidate in env_candidates:
        if candidate and Path(candidate).exists():
            return Path(candidate)

    for binary_name in ("soffice", "libreoffice"):
        binary_path = shutil.which(binary_name)
        if binary_path:
            return Path(binary_path)

    common_paths = [
        Path("/Applications/LibreOffice.app/Contents/MacOS/soffice"),
        Path("/usr/bin/soffice"),
        Path("/usr/local/bin/soffice"),
    ]
    for env_name in ("PROGRAMFILES", "PROGRAMFILES(X86)"):
        base_dir = os.environ.get(env_name)
        if base_dir:
            common_paths.append(Path(base_dir) / "LibreOffice" / "program" / "soffice.exe")

    for candidate in common_paths:
        if candidate.exists():
            return candidate
    return None
